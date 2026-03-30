#!/usr/bin/env python3
"""
split-docx.py — Split a .docx file into per-module markdown files.

Detects section boundaries by font size + bold formatting (since VisiGEO
requirement docs use formatting rather than heading styles), converts
the full document to GFM markdown via pandoc, splits at module boundaries,
and post-processes the markdown to fix common conversion issues.

Usage:
    python3 tools/split-docx.py <input.docx> [--output-dir docs/split]
    python3 tools/split-docx.py <input.docx> --level 1   # split at top-level only
    python3 tools/split-docx.py <input.docx> --level 2   # split at module level (default)

Requirements:
    - pandoc (brew install pandoc)
    - python-docx (pip install python-docx)
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


# ── Module name → English slug mapping ───────────────────────────────────

MODULE_SLUGS = {
    "产品概述": "product-overview",
    "通用规则": "common-rules",
    "全局导航": "global-nav",
    "首页": "landing",
    "网站GEO检测体验": "landing-geo-hook",
    "首页GEO检测体验": "landing-geo-hook",
    "邮箱密码登录": "auth-login",
    "登录/注册": "auth-login",
    "登录注册": "auth-login",
    "宣传模块": "landing-promo",
    "用户协议": "legal-pages",
    "隐私政策": "legal-pages",
    "用户协议&隐私政策": "legal-pages",
    "后台全局": "backend-global",
    "网站GEO检测": "geo-audit",
    "GEO内容创作": "content-creator",
    "LLMs文件生成": "llmstxt",
    "LLMs.txt文件生成": "llmstxt",
    "LLMs.txt": "llmstxt",
    "付费订阅体系": "subscription",
    "付费订阅": "subscription",
    "AI可见性分析": "ai-visibility",
    "AI可见性": "ai-visibility",
    "非功能需求": "non-functional",
    "数据需求": "non-functional",
    "SaaS前台": "saas-frontend",
    "SaaS后台功能": "saas-backend",
}

# Sections that should be further split into sub-sections
# Maps parent heading text → list of sub-heading texts to split on
# Defines how to expand container sections into their children for splitting.
# Key = heading text, Value = list of child heading texts to use as split points.
# Sections NOT listed here are kept as a single file (e.g., "1. 产品概述").
EXPAND_SECTIONS = {
    "3. SaaS前台": [
        "3.1 全局导航",
        "3.2.1 网站GEO检测体验",
        "3.2.2 邮箱密码登录/注册",
        "3.2.3 宣传模块",
        "3.3 用户协议&隐私政策页",
    ],
    "4. SaaS后台功能": [
        "4.1 后台全局",
        "4.2 网站GEO检测",
        "4.3 GEO内容创作",
        "4.4 LLMs.txt文件生成",
        "4.5 付费订阅体系",
        "4.6 AI可见性分析",
    ],
}


# ── Strikethrough detection via python-docx ──────────────────────────────

def detect_strikethrough_content(docx_path: str) -> list[dict]:
    """Detect paragraphs containing strikethrough text in the .docx.

    Returns list of dicts with paragraph index and strikethrough text.
    Strikethrough content typically means "deleted/deprecated" requirements
    that should NOT be included in the split output.
    """
    doc = Document(docx_path)
    results = []
    for idx, para in enumerate(doc.paragraphs):
        strike_runs = []
        normal_runs = []
        for r in para.runs:
            text = r.text.strip()
            if not text:
                continue
            if r.font.strike or r.font.double_strike:
                strike_runs.append(text)
            else:
                normal_runs.append(text)
        if strike_runs:
            results.append({
                "index": idx,
                "strike_text": " ".join(strike_runs),
                "normal_text": " ".join(normal_runs),
                "full_text": para.text.strip()[:120],
                "all_strike": len(normal_runs) == 0,  # entire paragraph is struck
            })
    return results


# ── Section detection via python-docx ────────────────────────────────────

def detect_all_headings(docx_path: str) -> list[dict]:
    """Parse the .docx and find all heading-like paragraphs by formatting.

    Returns headings sorted by document order with level/font_size info.
    """
    doc = Document(docx_path)
    headings = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph has bold runs
        bold_runs = [r for r in para.runs if r.text.strip() and r.bold]
        if not bold_runs:
            continue

        # Get font size from the first bold run with a size
        font_size = None
        for r in bold_runs:
            if r.font.size:
                font_size = r.font.size.pt
                break

        if font_size is None:
            continue

        # Determine heading level from numbering pattern and font size
        level = None
        if re.match(r"^\d+\.\s+\S", text) and font_size >= 15.5:
            level = 1  # e.g., "1. 产品概述" at 16pt
        elif re.match(r"^\d+\.\d+\s+\S", text) and font_size >= 14.5:
            level = 2  # e.g., "3.1 全局导航" at 15pt
        elif re.match(r"^\d+\.\d+\.\d+\s+\S", text) and font_size >= 13.5:
            level = 3  # e.g., "3.2.1 网站GEO检测体验" at 14pt
        elif re.match(r"^\d+\.\d+\.\d+\.\d+\s+\S", text) and font_size >= 11.5:
            level = 4  # e.g., "3.2.1.1 界面元素" at 12pt

        if level is not None:
            headings.append({
                "index": idx,
                "text": text,
                "level": level,
                "font_size": font_size,
            })

    return headings


def select_split_headings(headings: list[dict], split_level: int) -> list[dict]:
    """Select which headings to use as split points.

    For split_level=2 (default, module-level), we produce a ~14-file
    structure matching the existing docs/split/:
      - Level 1 headings NOT in EXPAND_SECTIONS → single file (keep as-is)
      - Level 1 headings IN EXPAND_SECTIONS → replaced by their children
    For split_level=1, only top-level headings are used (5 files).
    """
    if split_level == 1:
        return [h for h in headings if h["level"] == 1]

    # Collect the child heading texts we want from all EXPAND entries
    all_children = set()
    for children in EXPAND_SECTIONS.values():
        all_children.update(children)

    selected = []
    for h in headings:
        text = h["text"]
        level = h["level"]

        # Level 1: include only if it's NOT a container to be expanded
        if level == 1:
            if text not in EXPAND_SECTIONS:
                selected.append(h)
            # else: skip — its children will be picked up below

        # Any level: include if it's listed as a child in EXPAND_SECTIONS
        elif text in all_children:
            selected.append(h)

    # Sort by document order
    selected.sort(key=lambda x: x["index"])
    return selected


# ── Filename generation ──────────────────────────────────────────────────

def slugify(title: str) -> str:
    """Generate a filename-safe English slug from a section title."""
    # Remove numbering prefix
    core = re.sub(r"^[\d.]+\s*", "", title).strip()
    # Try to match against known module slugs
    for cn_key, en_slug in MODULE_SLUGS.items():
        if cn_key in core or cn_key in title:
            return en_slug
    # Fallback: keep alphanumeric chars
    slug = core.lower()
    slug = re.sub(r"[^\w\s-]", "", slug)
    slug = re.sub(r"[\s_]+", "-", slug).strip("-")
    return slug[:60] if slug else "section"


def generate_filename(index: int, title: str) -> str:
    """Generate a numbered filename like 00-产品概述-product-overview.md."""
    cn_name = re.sub(r"^[\d.]+\s*", "", title).strip()
    # Clean up common suffixes and unsafe filename chars
    cn_name = re.sub(r"\s*[（(].*?[）)]", "", cn_name).strip()
    cn_name = cn_name.rstrip("页")  # "隐私政策页" → "隐私政策"
    cn_name = cn_name.replace("/", "").replace("\\", "")  # remove path separators
    cn_name = re.sub(r"[&]+", "", cn_name).strip()  # remove & characters
    slug = slugify(title)
    return f"{index:02d}-{cn_name}-{slug}.md"


# ── Pandoc conversion ────────────────────────────────────────────────────

def convert_docx_to_markdown(docx_path: str, media_dir: str) -> str:
    """Convert .docx to GFM markdown using pandoc, extracting images."""
    cmd = [
        "pandoc",
        docx_path,
        "--to=gfm",
        "--wrap=none",
        f"--extract-media={media_dir}",
        "--standalone",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
    return result.stdout


# ── Markdown splitting ───────────────────────────────────────────────────

def build_section_regex(title: str) -> str:
    """Build a regex that matches a section heading in pandoc's markdown output.

    Pandoc converts bold paragraphs (not real headings) in various ways:
      - **1. 产品概述**
      - 1\\. **产品概述**
      - **1.** **产品概述**
    """
    core = re.sub(r"^[\d.]+\s*", "", title).strip()
    escaped_core = re.escape(core)
    # Extract the numbering prefix (e.g., "3.2.1")
    num_match = re.match(r"^([\d.]+)", title)
    num_prefix = num_match.group(1) if num_match else ""
    escaped_num = re.escape(num_prefix).replace(r"\.", r"\\?\.")

    variants = [
        # # N. Title or ## N. Title
        rf"^#{1,3}\s+{escaped_num}\s+{escaped_core}",
        # **N. Title**
        rf"^\*\*{escaped_num}\s+{escaped_core}\*\*",
        # N\. **Title**
        rf"^{escaped_num}\s+\*\*{escaped_core}\*\*",
        # **N.** **Title**
        rf"^\*\*{escaped_num}\*\*\s+\*\*{escaped_core}\*\*",
        # Plain: N. Title (without bold, but exact match)
        rf"^{escaped_num}\s+{escaped_core}$",
    ]
    return "|".join(f"(?:{v})" for v in variants)


def split_markdown_by_sections(
    md_content: str, sections: list[dict]
) -> list[tuple[str, str]]:
    """Split markdown content at each section heading.

    Returns list of (title, content) tuples.
    """
    if not sections:
        return [("full-document", md_content)]

    # Find each section heading's position in the markdown
    split_points = []
    for sec in sections:
        pattern = build_section_regex(sec["text"])
        match = re.search(pattern, md_content, re.MULTILINE)
        if match:
            split_points.append((match.start(), sec["text"]))
        else:
            print(f"   WARNING: Could not find heading '{sec['text']}' in markdown")

    split_points.sort(key=lambda x: x[0])

    if not split_points:
        return [("full-document", md_content)]

    # Extract content between split points
    result = []
    preamble_end = split_points[0][0]

    for i, (pos, title) in enumerate(split_points):
        end = split_points[i + 1][0] if i + 1 < len(split_points) else len(md_content)
        content = md_content[pos:end].strip()
        result.append((title, content))

    # Prepend preamble (doc title, metadata) to first section
    if preamble_end > 50:
        preamble = md_content[:preamble_end].strip()
        first_title, first_content = result[0]
        result[0] = (first_title, preamble + "\n\n" + first_content)

    return result


# ── Markdown post-processing ─────────────────────────────────────────────

def convert_single_cell_tables_to_codeblocks(content: str) -> str:
    """Convert single-cell HTML tables (1 row, 1 col) to markdown code blocks.

    Chinese docs commonly use single-cell tables as "code containers" to display
    prompt templates, SQL, config, etc.  Pandoc faithfully converts these to
    HTML <table> blocks, but we want clean markdown code blocks instead.

    Pattern detected:
      <table>
      <colgroup><col style="width: 100%" /></colgroup>
      <tbody><tr>
      <td ...>CONTENT with <br /> line breaks and &lt;entity&gt; encoding</td>
      </tr></tbody></table>
    """
    # Match single-cell tables (width: 100%, one <td>)
    pattern = re.compile(
        r"<table>\s*\n"
        r"<colgroup>\s*\n"
        r'<col\s+style="width:\s*100%"\s*/>\s*\n'
        r"</colgroup>\s*\n"
        r"<tbody>\s*\n"
        r"<tr>\s*\n"
        r'<td[^>]*>(.*?)</td>\s*\n'
        r"</tr>\s*\n"
        r"</tbody>\s*\n"
        r"</table>",
        re.DOTALL,
    )

    def replace_with_codeblock(m: re.Match) -> str:
        cell_content = m.group(1)
        # Decode HTML entities
        cell_content = cell_content.replace("&lt;", "<")
        cell_content = cell_content.replace("&gt;", ">")
        cell_content = cell_content.replace("&amp;", "&")
        cell_content = cell_content.replace("&quot;", '"')
        cell_content = cell_content.replace("&#39;", "'")
        # Convert <br /> to newlines
        cell_content = re.sub(r"<br\s*/?>", "\n", cell_content)
        # Remove remaining HTML tags (<p>, <strong>, etc.)
        cell_content = re.sub(r"</?(?:p|strong|em|span|div)[^>]*>", "", cell_content)
        # Clean up whitespace
        cell_content = cell_content.strip()
        # Detect language hint from first line (e.g., "SQL", "XML", "JSON", "YAML")
        first_line = cell_content.split("\n")[0].strip()
        lang = ""
        if first_line in ("SQL", "XML", "JSON", "YAML", "HTML", "CSS", "JS"):
            lang = first_line.lower()
            # Remove the language tag line from content
            cell_content = "\n".join(cell_content.split("\n")[1:]).strip()
        elif first_line.startswith("```"):
            # Already has a code fence, skip
            return cell_content

        return f"```{lang}\n{cell_content}\n```"

    return pattern.sub(replace_with_codeblock, content)


def remove_strikethrough(content: str) -> str:
    """Remove strikethrough text (~~text~~) from markdown.

    Strikethrough in requirements docs indicates deleted/deprecated content
    that should not be included in the output.  Handles:
      - Entire lines that are struck through → remove the line
      - Inline struck text → remove just that portion
    """
    # Remove lines that are entirely strikethrough (possibly with whitespace)
    content = re.sub(r"^\s*~~[^~]+~~\s*$", "", content, flags=re.MULTILINE)
    # Remove inline strikethrough spans
    content = re.sub(r"~~([^~]+)~~", "", content)
    return content


def fix_markdown(content: str, media_dir: str, strip_strikethrough: bool = True) -> str:
    """Fix common pandoc markdown conversion issues."""
    # First pass: convert single-cell tables to code blocks
    content = convert_single_cell_tables_to_codeblocks(content)

    # Remove strikethrough content if requested
    if strip_strikethrough:
        content = remove_strikethrough(content)

    lines = content.split("\n")
    fixed = []

    for line in lines:
        # Fix escaped numbered lists: "1\." → "1."
        line = re.sub(r"^(\s*)(\d+)\\(\.\s)", r"\1\2\3", line)

        # Fix image paths in markdown syntax: ![alt](path) → ![alt](media/filename)
        line = re.sub(
            r"!\[([^\]]*)\]\([^)]*?/media/([^)/]+)\)",
            r"![\1](media/\2)",
            line,
        )

        # Fix image paths in HTML <img> tags: src="..." → src="media/filename"
        line = re.sub(
            r'(<img\s[^>]*?)src="[^"]*?/media/([^"/]+)"',
            r'\1src="media/\2"',
            line,
        )

        # Convert HTML <img> tags to markdown syntax for cleaner output
        img_match = re.match(
            r'^<img\s[^>]*?src="([^"]+)"[^>]*?/?>$', line.strip()
        )
        if img_match:
            img_src = img_match.group(1)
            line = f"![]({img_src})"

        # Fix escaped brackets: \[text\] → [text]
        line = line.replace("\\[", "[").replace("\\]", "]")

        # Fix stray HTML line breaks
        line = re.sub(r"<br\s*/?>", "  ", line)

        # Fix non-breaking spaces
        line = line.replace("\u00a0", " ")

        # Fix zero-width spaces and other invisible chars
        line = line.replace("\u200b", "")
        line = line.replace("\ufeff", "")

        fixed.append(line)

    result = "\n".join(fixed)

    # Normalize excessive blank lines (max 2 consecutive)
    result = re.sub(r"\n{4,}", "\n\n\n", result)

    # Fix pandoc dash-separated pseudo-tables
    result = re.sub(
        r"^\s*-{3,}\s+-{3,}\s*$",
        "",
        result,
        flags=re.MULTILINE,
    )

    # Convert bold-only paragraphs that look like headings into proper markdown headings.
    # Patterns: "**N. Title**", "N. **Title**", "N.N **Title**"
    # Level 1: N.
    result = re.sub(
        r"^\*\*(\d+\.)\s+(.+?)\*\*$",
        r"# \1 \2",
        result,
        flags=re.MULTILINE,
    )
    result = re.sub(
        r"^(\d+\.)\s+\*\*(.+?)\*\*$",
        r"# \1 \2",
        result,
        flags=re.MULTILINE,
    )
    # Level 2: N.N
    result = re.sub(
        r"^\*\*(\d+\.\d+)\s+(.+?)\*\*$",
        r"## \1 \2",
        result,
        flags=re.MULTILINE,
    )
    result = re.sub(
        r"^(\d+\.\d+)\s+\*\*(.+?)\*\*$",
        r"## \1 \2",
        result,
        flags=re.MULTILINE,
    )
    # Level 3: N.N.N
    result = re.sub(
        r"^\*\*(\d+\.\d+\.\d+)\s+(.+?)\*\*$",
        r"### \1 \2",
        result,
        flags=re.MULTILINE,
    )
    result = re.sub(
        r"^(\d+\.\d+\.\d+)\s+\*\*(.+?)\*\*$",
        r"### \1 \2",
        result,
        flags=re.MULTILINE,
    )
    # Level 4: N.N.N.N
    result = re.sub(
        r"^\*\*(\d+\.\d+\.\d+\.\d+)\s+(.+?)\*\*$",
        r"#### \1 \2",
        result,
        flags=re.MULTILINE,
    )
    result = re.sub(
        r"^(\d+\.\d+\.\d+\.\d+)\s+\*\*(.+?)\*\*$",
        r"#### \1 \2",
        result,
        flags=re.MULTILINE,
    )

    # Ensure file ends with single newline
    result = result.strip() + "\n"

    return result


# ── Image handling ───────────────────────────────────────────────────────

def copy_media_files(pandoc_media_dir: str, output_media_dir: str) -> int:
    """Copy extracted media files to the output directory."""
    if not os.path.exists(pandoc_media_dir):
        return 0

    os.makedirs(output_media_dir, exist_ok=True)
    count = 0
    for root, _dirs, files in os.walk(pandoc_media_dir):
        for f in files:
            src = os.path.join(root, f)
            dst = os.path.join(output_media_dir, f)
            if os.path.exists(dst):
                name, ext = os.path.splitext(f)
                i = 1
                while os.path.exists(dst):
                    dst = os.path.join(output_media_dir, f"{name}-{i}{ext}")
                    i += 1
            shutil.copy2(src, dst)
            count += 1
    return count


# ── README generation ────────────────────────────────────────────────────

def generate_readme(
    file_info: list[tuple[str, str, str]], source_file: str
) -> str:
    """Generate README.md index.  file_info: list of (filename, title, content)."""
    source_name = os.path.basename(source_file)
    lines = [
        "# VisiGEO Requirements - Split by Module",
        "",
        f"Source: `docs/req/{source_name}`",
        "",
        "## Module Index",
        "",
        "| # | File | Module | Lines |",
        "|---|------|--------|-------|",
    ]

    for i, (filename, title, content) in enumerate(file_info):
        line_count = len(content.strip().split("\n"))
        clean_title = re.sub(r"^[\d.]+\s*", "", title).strip()
        lines.append(
            f"| {i:02d} | [{filename}]({filename}) | {clean_title} | {line_count} |"
        )

    lines.append("")
    return "\n".join(lines)


# ── Main ─────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Split a .docx file into per-module markdown files"
    )
    parser.add_argument("input", help="Path to the .docx file")
    parser.add_argument(
        "--output-dir",
        default="docs/split",
        help="Output directory (default: docs/split)",
    )
    parser.add_argument(
        "--level",
        type=int,
        default=2,
        choices=[1, 2],
        help="Split level: 1=top-level only, 2=module level (default: 2)",
    )
    parser.add_argument(
        "--clean",
        action="store_true",
        help="Remove existing .md and media/ in output directory first",
    )
    parser.add_argument(
        "--keep-strikethrough",
        action="store_true",
        help="Keep strikethrough content instead of removing it",
    )
    args = parser.parse_args()

    input_path = os.path.abspath(args.input)
    output_dir = os.path.abspath(args.output_dir)

    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    print(f"Input:  {input_path}")
    print(f"Output: {output_dir}")
    print(f"Level:  {args.level}")

    strip_strike = not args.keep_strikethrough

    # Step 0: Detect strikethrough content
    strike_items = detect_strikethrough_content(input_path)
    if strike_items:
        print(f"\n[0/6] Found {len(strike_items)} paragraphs with strikethrough content:")
        for s in strike_items[:10]:
            marker = "[FULL]" if s["all_strike"] else "[PARTIAL]"
            print(f"   {marker} {s['full_text']}")
        if len(strike_items) > 10:
            print(f"   ... and {len(strike_items) - 10} more")
        if strip_strike:
            print("   -> Strikethrough content will be REMOVED from output")
        else:
            print("   -> Strikethrough content will be KEPT (--keep-strikethrough)")
    else:
        print("\n[0/6] No strikethrough content detected")

    # Step 1: Detect headings from .docx formatting
    print("\n[1/6] Detecting headings from .docx formatting...")
    all_headings = detect_all_headings(input_path)
    print(f"   Found {len(all_headings)} headings total")
    for h in all_headings:
        indent = "  " * (h["level"] - 1)
        print(f"   {indent}L{h['level']} ({h['font_size']}pt) {h['text']}")

    # Step 2: Select split points
    print(f"\n[2/6] Selecting split points (level={args.level})...")
    split_headings = select_split_headings(all_headings, args.level)
    print(f"   Will split into {len(split_headings)} sections:")
    for i, h in enumerate(split_headings):
        print(f"   {i:2d}. {h['text']}")

    # Step 3: Convert to markdown with pandoc
    print("\n[3/6] Converting to markdown with pandoc...")
    with tempfile.TemporaryDirectory() as tmp_dir:
        media_dir = os.path.join(tmp_dir, "media")
        md_content = convert_docx_to_markdown(input_path, media_dir)
        print(f"   Markdown: {len(md_content)} chars, {len(md_content.splitlines())} lines")

        # Step 4: Split by sections
        print("\n[4/6] Splitting markdown by sections...")
        split_result = split_markdown_by_sections(md_content, split_headings)
        print(f"   Got {len(split_result)} parts")

        # Prepare output directory
        if args.clean and os.path.exists(output_dir):
            for f in os.listdir(output_dir):
                fp = os.path.join(output_dir, f)
                if f.endswith(".md"):
                    os.remove(fp)
                elif f == "media" and os.path.isdir(fp):
                    shutil.rmtree(fp)
        os.makedirs(output_dir, exist_ok=True)

        # Step 5: Copy media files
        print("\n[5/6] Copying media files...")
        output_media = os.path.join(output_dir, "media")
        # Pandoc puts media under <extract-media>/media/
        pandoc_media = os.path.join(media_dir, "media")
        img_count = copy_media_files(pandoc_media, output_media)
        if img_count == 0:
            img_count = copy_media_files(media_dir, output_media)
        print(f"   Copied {img_count} images")

        # Step 6: Write split files + README
        print("\n[6/6] Writing split markdown files...")
        file_info = []
        for i, (title, content) in enumerate(split_result):
            filename = generate_filename(i, title)
            fixed = fix_markdown(content, media_dir, strip_strikethrough=strip_strike)

            filepath = os.path.join(output_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(fixed)

            lc = len(fixed.strip().split("\n"))
            print(f"   {filename} ({lc} lines)")
            file_info.append((filename, title, fixed))

        readme = generate_readme(file_info, input_path)
        with open(os.path.join(output_dir, "README.md"), "w", encoding="utf-8") as f:
            f.write(readme)
        print("   README.md")

    print(f"\nDone! {len(split_result)} files written to {output_dir}/")
    return 0


if __name__ == "__main__":
    sys.exit(main())
